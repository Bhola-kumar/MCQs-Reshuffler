import random
import docx
import os

Q_list=[]
O_list=[]
copy_list=[]
O_O_list=[]

def option_reshuffler(Op_list,option_container_list):
 random.shuffle(Op_list)
 copy_list=Op_list.copy()
 option_container_list.append(copy_list)
 Op_list.clear()

no_of_O=int(input("Number of Option per Question: "))
div=no_of_O+1
file_name=input("Enter the Doc file name: ")
doc=docx.Document('D://OneDrive//Documents//'+file_name+'.docx')
no_of_Q=int((len(doc.paragraphs))/(no_of_O+1))
print("Total number of Question present in the Given DOC is:",no_of_Q)

for i in range(len(doc.paragraphs)):
    if i%div==0:
        Q_list.append(doc.paragraphs[i].text)
    else:
        O_list.append(doc.paragraphs[i].text)
        if len(O_list)==no_of_O:
            option_reshuffler(O_list,O_O_list)

final_shuffle=list(zip(Q_list,O_O_list))
random.shuffle(final_shuffle)
Q_list,O_O_list = zip(*final_shuffle)

output_doc=docx.Document()

for p in range(no_of_Q):
 output_doc.add_paragraph('Q'+str(p+1)+'. ' + Q_list[p])
 for q in range(no_of_O):
  output_doc.add_paragraph(chr(q+65)+'. ' + O_O_list[p][q])

print("Questions have been Shuffled successfully!!")
output_file_name=input("Enter the new Doc file name for output: ")
output_doc.save('D://OneDrive//Documents//'+output_file_name+'.docx')
os.system('start  D://OneDrive//Documents/'+output_file_name+'.docx')